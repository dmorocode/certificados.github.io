import pandas as pd
from docx import Document
# Remova esta linha se ainda estiver aqui no seu arquivo:
# from docx2pdf import convert
from PyPDF2 import PdfMerger
import os
import re
import shutil
import traceback
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
from flask_cors import CORS
import time
import zipfile
import json
import hashlib
from functools import wraps
import subprocess # Módulo para rodar comandos externos como pandoc

app = Flask(__name__, template_folder='templates')
CORS(app,
     supports_credentials=True,
     origins=["http://127.0.0.1:5000", "http://localhost:5000"]
)

# --- Configurações de Pastas ---
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'temp_uploads_web')
GENERATED_CERTIFICATES_FOLDER = os.path.join(BASE_DIR, 'generated_certificates_web')
USERS_DB_PATH = os.path.join(BASE_DIR, 'users_db.json')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_CERTIFICATES_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GENERATED_CERTIFICATES_FOLDER'] = GENERATED_CERTIFICATES_FOLDER
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=60)

def load_users_db():
    if os.path.exists(USERS_DB_PATH):
        with open(USERS_DB_PATH, 'r') as f:
            return json.load(f)
    return {}

def save_users_db(db):
    with open(USERS_DB_PATH, 'w') as f:
        json.dump(db, f, indent=4)

users_db = load_users_db()

if "danilo" not in users_db:
    danilo_password_hash = hashlib.sha256("@danilo30!".encode('utf-8')).hexdigest()
    users_db["danilo"] = {
        "password_hash": danilo_password_hash,
        "role": "admin"
    }
    save_users_db(users_db)
    print("INFO: Usuário 'danilo' (admin) padrão criado/verificado com senha '@danilo30!' (hashed).")

def require_auth(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            print("INFO: Autenticação falhou: Cabeçalho 'Authorization' ausente ou mal formatado.")
            return jsonify({"error": "Autenticação requerida. Token de acesso ausente ou inválido."}), 401

        token = auth_header.split(' ')[1]

        if token in users_db:
            request.current_user = token
            request.current_user_role = users_db[token].get("role", "user")
            return f(*args, **kwargs)
        else:
            print(f"INFO: Autenticação falhou: Token '{token}' inválido.")
            return jsonify({"error": "Token de acesso inválido ou expirado."}), 401
    return decorated_function

def require_role(role):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not hasattr(request, 'current_user') or not hasattr(request, 'current_user_role'):
                print("ERRO: Erro interno de autenticação - usuário ou role não definidos após require_auth.")
                return jsonify({"error": "Erro de autenticação interno: usuário não definido."}), 500

            if request.current_user_role == role:
                return f(*args, **kwargs)
            else:
                print(f"INFO: Acesso negado para '{request.current_user}'. Requer role: '{role}', mas é '{request.current_user_role}'.")
                return jsonify({"error": f"Acesso negado. Requer role: '{role}'."}), 403
        return decorated_function
    return decorator

@app.route('/')
def index():
    return render_template('gerador_web.html')

def process_document_tags(document_obj, data_row):
    for paragraph in document_obj.paragraphs:
        full_paragraph_text = paragraph.text
        if '{' in full_paragraph_text and '}' in full_paragraph_text:
            original_runs = list(paragraph.runs)
            replaced_text = replace_all_tags_in_text(full_paragraph_text, data_row)
            if replaced_text != full_paragraph_text:
                apply_formatting_runs(paragraph, original_runs, replaced_text)

    for table in document_obj.tables:
        for row_table in table.rows:
            for cell in row_table.cells:
                for paragraph_in_cell in cell.paragraphs:
                    full_cell_text = paragraph_in_cell.text
                    if '{' in full_cell_text and '}' in full_cell_text:
                        original_runs_cell = list(paragraph_in_cell.runs)
                        replaced_cell_text = replace_all_tags_in_text(full_cell_text, data_row)
                        if replaced_cell_text != full_cell_text:
                            apply_formatting_runs(paragraph_in_cell, original_runs_cell, replaced_cell_text)

def replace_all_tags_in_text(text_content, data_row_normalized):
    tags_found_in_word = re.findall(r'\{(.*?)\}', text_content)
    for tag_content in tags_found_in_word:
        normalized_tag_from_word = tag_content.strip().lower()
        replacement_value = str(data_row_normalized.get(normalized_tag_from_word, f'{{{tag_content}}}'))
        text_content = text_content.replace(f'{{{tag_content}}}', replacement_value)
    return text_content

def apply_formatting_runs(paragraph, original_runs, replaced_text):
    paragraph.clear()
    if original_runs:
        first_run_original = original_runs[0]
        new_run = paragraph.add_run(replaced_text)
        new_run.bold = first_run_original.bold
        new_run.italic = first_run_original.italic
        new_run.underline = first_run_original.underline
        if first_run_original.font.name:
            new_run.font.name = first_run_original.font.name
        if first_run_original.font.size:
            new_run.font.size = first_run_original.font.size
        if first_run_original.font.color and first_run_original.font.color.rgb:
            new_run.font.color.rgb = first_run_original.font.color.rgb
    else:
        paragraph.add_run(replaced_text)

def create_zip_from_folder(folder_path, zip_path):
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(folder_path):
            for file_item in files:
                file_path_item = os.path.join(root, file_item)
                archive_name = os.path.relpath(file_path_item, folder_path)
                zipf.write(file_path_item, archive_name)
    print(f"INFO: Ficheiro ZIP criado em: {zip_path}")

@app.route('/api/users', methods=['GET'])
@require_auth
@require_role('admin')
def get_users_api():
    user_list = []
    for user, data in users_db.items():
        user_list.append({
            "username": user,
            "role": data.get("role", "user")
        })
    return jsonify(user_list), 200

@app.route('/api/users/create', methods=['POST'])
@require_auth
@require_role('admin')
def create_user_api():
    data = request.get_json()
    if not data or not data.get('username') or not data.get('password'):
        return jsonify({"error": "Nome de utilizador e senha são obrigatórios."}), 400

    username = data['username'].strip()
    password = data['password']
    role = data.get('role', 'user').strip().lower()

    if not username or not password:
        return jsonify({"error": "Nome de utilizador e senha não podem ser vazios."}), 400
    if role not in ['user', 'admin']:
        return jsonify({"error": "Role inválida. Use 'user' ou 'admin'."}), 400

    if username in users_db:
        return jsonify({"error": f"Utilizador '{username}' já existe."}), 409

    password_hash = hashlib.sha256(password.encode('utf-8')).hexdigest()

    users_db[username] = {
        "password_hash": password_hash,
        "role": role
    }
    save_users_db(users_db)
    print(f"INFO: Utilizador '{username}' (role: {role}) criado (em memória e persistido).")
    return jsonify({"message": f"Utilizador '{username}' criado com sucesso.", "username": username, "role": role}), 201

@app.route('/api/login', methods=['POST'])
def login_api():
    data = request.get_json()
    if not data or not data.get('username') or not data.get('password'):
        return jsonify({"error": "Nome de utilizador e senha são obrigatórios."}), 400

    username = data['username']
    password_input = data['password']

    user_data = users_db.get(username)

    if user_data:
        input_password_hash = hashlib.sha256(password_input.encode('utf-8')).hexdigest()
        if user_data['password_hash'] == input_password_hash:
            print(f"INFO: Tentativa de login bem-sucedida para '{username}'.")
            return jsonify(
                message="Login bem-sucedido.",
                user=username,
                role=user_data.get("role", "user"),
                token=username
            ), 200

    print(f"INFO: Tentativa de login falhada para o utilizador '{username}'.")
    return jsonify({"error": "Nome de utilizador ou senha inválidos."}), 401

@app.route('/generate_certificates_api', methods=['POST'])
@require_auth
def generate_certificates_api():
    print(f"INFO: Endpoint /generate_certificates_api acedido por '{request.current_user}'.")

    if 'excelFile' not in request.files or 'templateFile' not in request.files:
        return jsonify({"error": "Arquivos Excel e Template são obrigatórios."}), 400

    excel_file = request.files['excelFile']
    template_file = request.files['templateFile']

    if excel_file.filename == '' or template_file.filename == '':
        return jsonify({"error": "Nomes de arquivo não podem ser vazios."}), 400

    request_id_suffix = request.current_user
    request_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f") + f"_{request_id_suffix}"

    current_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], request_id)
    session_output_path = os.path.join(app.config['GENERATED_CERTIFICATES_FOLDER'], f"output_{request_id}")
    zip_staging_folder = os.path.join(session_output_path, "staging_for_zip")
    individual_docx_staging_folder = os.path.join(zip_staging_folder, "WORDs_Individuais")
    individual_pdf_staging_folder = os.path.join(zip_staging_folder, "PDFs_Individuais")

    os.makedirs(current_upload_folder, exist_ok=True)
    os.makedirs(session_output_path, exist_ok=True)
    os.makedirs(zip_staging_folder, exist_ok=True)
    os.makedirs(individual_docx_staging_folder, exist_ok=True)
    os.makedirs(individual_pdf_staging_folder, exist_ok=True)

    excel_path = os.path.join(current_upload_folder, secure_filename(excel_file.filename))
    template_path = os.path.join(current_upload_folder, secure_filename(template_file.filename))

    try:
        excel_file.save(excel_path)
        template_file.save(template_path)
    except Exception as e:
        print(f"ERRO: Falha ao salvar arquivos de upload: {e}")
        traceback.print_exc()
        shutil.rmtree(current_upload_folder, ignore_errors=True)
        return jsonify({"error": f"Erro ao salvar arquivos: {e}"}), 500

    any_processing_errors = False

    try:
        df = pd.read_excel(excel_path)
    except Exception as e:
        traceback.print_exc()
        shutil.rmtree(current_upload_folder, ignore_errors=True)
        shutil.rmtree(session_output_path, ignore_errors=True)
        return jsonify({"error": f"Erro ao ler arquivo Excel. Verifique o formato do arquivo: {e}"}), 500

    total_participants_overall = df.apply(lambda x: x.dropna().count()).sum()
    if total_participants_overall == 0:
        shutil.rmtree(current_upload_folder, ignore_errors=True)
        shutil.rmtree(session_output_path, ignore_errors=True)
        return jsonify({"message": "Nenhum participante encontrado no arquivo Excel.", "zip_file_path": None}), 200

    processed_participants_overall_count = 0

    for col_idx, column_excel_header in enumerate(df.columns):
        current_col_header_str = str(column_excel_header)
        clean_column_name_for_folder = re.sub(r'[^\w\s-]', '', current_col_header_str).strip().replace(' ', '_')
        if not clean_column_name_for_folder:
            clean_column_name_for_folder = f"Coluna_{col_idx + 1}"

        column_aggregate_staging_folder = os.path.join(zip_staging_folder, clean_column_name_for_folder)
        os.makedirs(column_aggregate_staging_folder, exist_ok=True)

        participant_names_in_column = df[column_excel_header].dropna()
        if participant_names_in_column.empty:
            continue

        pdfs_for_this_column_merger_paths = []

        for item_idx, participant_name_raw in enumerate(participant_names_in_column):
            participant_name = str(participant_name_raw).strip()
            if not participant_name:
                continue
            processed_participants_overall_count += 1

            data_for_tags = {'nome': participant_name}

            base_filename_participant = re.sub(r'[^\w\s-]', '', participant_name).strip().replace(' ', '_').replace('/', '_')
            if not base_filename_participant:
                base_filename_participant = f"Participante_{col_idx + 1}_{item_idx + 1}"
            else:
                base_filename_participant = f"{base_filename_participant}_{col_idx +1}_{item_idx + 1}"

            output_docx_name = os.path.join(individual_docx_staging_folder, f"certificado_{base_filename_participant}.docx")
            output_pdf_name = os.path.join(individual_pdf_staging_folder, f"certificado_{base_filename_participant}.pdf")

            try:
                doc = Document(template_path)
                process_document_tags(doc, data_for_tags)
                doc.save(output_docx_name)
                time.sleep(0.1)

                # --- Início do Bloco de Conversão DOCX para PDF (Pandoc) ---
                try:
                    print(f"INFO: Tentando converter {output_docx_name} para {output_pdf_name} usando Pandoc.")

                    # Garante que o diretório de saída para o PDF exista
                    os.makedirs(os.path.dirname(output_pdf_name), exist_ok=True)

                    # Comando para Pandoc
                    # -o: arquivo de saída
                    # --from docx: formato de entrada é DOCX
                    # --to pdf: formato de saída é PDF
                    # --pdf-engine=xelatex: usa xelatex para gerar PDF (melhor para UTF-8)
                    pandoc_command = [
                        'pandoc',
                        output_docx_name,
                        '-o', output_pdf_name,
                        '--from', 'docx',
                        '--to', 'pdf',
                        '--pdf-engine=xelatex' # xelatex é geralmente melhor para caracteres especiais e fontes
                    ]
                    print(f"INFO: Comando pandoc: {' '.join(pandoc_command)}")

                    result = subprocess.run(
                        pandoc_command,
                        check=True, # Lança CalledProcessError se o comando falhar
                        capture_output=True, # Captura stdout e stderr
                        text=True, # Decodifica a saída como texto
                        timeout=180 # Aumenta o timeout para 180 segundos (3 minutos)
                    )
                    print(f"INFO: Pandoc STDOUT para {participant_name}: {result.stdout}")
                    print(f"INFO: Pandoc STDERR para {participant_name}: {result.stderr}")

                    time.sleep(0.05)
                    if os.path.exists(output_pdf_name):
                        pdfs_for_this_column_merger_paths.append(output_pdf_name)
                        print(f"SUCESSO: PDF gerado para {participant_name} em {output_pdf_name}")
                    else:
                        error_details = f"Pandoc falhou em criar o PDF em {output_pdf_name}."
                        if result.stdout: error_details += f"\nSTDOUT: {result.stdout}"
                        if result.stderr: error_details += f"\nSTDERR: {result.stderr}"
                        raise Exception(error_details)

                except subprocess.CalledProcessError as sub_error:
                    error_msg = f"ERRO Pandoc para {participant_name}: Comando falhou com código {sub_error.returncode}."
                    error_msg += f"\nSTDOUT: {sub_error.stdout}"
                    error_msg += f"\nSTDERR: {sub_error.stderr}"
                    print(error_msg)
                    traceback.print_exc()
                    any_processing_errors = True
                except subprocess.TimeoutExpired:
                    print(f"ERRO: Conversão de PDF para {participant_name} excedeu o tempo limite de 180 segundos (Pandoc).")
                    any_processing_errors = True
                except Exception as pdf_conversion_error:
                    print(f"ERRO geral na conversão PDF (Pandoc) para {participant_name}: {pdf_conversion_error}")
                    traceback.print_exc()
                    any_processing_errors = True
                # --- Fim do Bloco de Conversão DOCX para PDF (Pandoc) ---

            except Exception as e_docx:
                print(f"ERRO DOCX para {participant_name}: {e_docx}")
                traceback.print_exc()
                any_processing_errors = True
                continue

        if pdfs_for_this_column_merger_paths:
            merger = PdfMerger()
            for pdf_path in pdfs_for_this_column_merger_paths:
                try:
                    if os.path.exists(pdf_path):
                        merger.append(pdf_path)
                except Exception as merge_err:
                    print(f"ERRO ao adicionar PDF '{pdf_path}' ao agregador da coluna: {merge_err}")
                    traceback.print_exc()
                    any_processing_errors = True

            if merger.inputs:
                merged_filename = f"TODOS_CERTIFICADOS_{clean_column_name_for_folder}.pdf"
                final_column_merged_path = os.path.join(column_aggregate_staging_folder, merged_filename)
                try:
                    merger.write(final_column_merged_path)
                    merger.close()
                except Exception as e_merge:
                    print(f"ERRO ao salvar PDF agregado da coluna '{clean_column_name_for_folder}': {e_merge}")
                    traceback.print_exc()
                    any_processing_errors = True
            else:
                merger.close()

    zip_filename = "Certificados_Gerados.zip"
    zip_file_full_path = os.path.join(session_output_path, zip_filename)

    try:
        create_zip_from_folder(zip_staging_folder, zip_file_full_path)
        shutil.rmtree(zip_staging_folder, ignore_errors=True)
    except Exception as e_zip:
        print(f"ERRO ao criar ou limpar o ficheiro ZIP/pasta staging: {e_zip}")
        traceback.print_exc()
        any_processing_errors = True
        return jsonify({"error": f"Erro ao criar ficheiro ZIP: {e_zip}", "zip_file_path": None}), 500

    try:
        shutil.rmtree(current_upload_folder, ignore_errors=True)
    except Exception as e_clean:
        print(f"ERRO ao remover pasta de upload temporária '{current_upload_folder}': {e_clean}")
        traceback.print_exc()

    final_message = "Certificados processados e empacotados com sucesso."
    if any_processing_errors:
        final_message += " No entanto, alguns erros ocorreram durante o processo de geração/conversão de PDFs. Verifique os logs do servidor para mais detalhes."

    zip_download_path = os.path.join(os.path.basename(session_output_path), zip_filename)
    print(f"INFO: {final_message}. Caminho do ZIP para download: {zip_download_path}")
    return jsonify({"message": final_message, "zip_file_path": zip_download_path}), 200

@app.route('/download_generated_file/<path:filepath>')
@require_auth
def download_generated_file(filepath):
    print(f"INFO: Pedido de download para: '{filepath}' por '{request.current_user}'.")

    full_file_path = os.path.join(app.config['GENERATED_CERTIFICATES_FOLDER'], filepath)
    resolved_file_path = os.path.realpath(full_file_path)

    allowed_dir = os.path.realpath(app.config['GENERATED_CERTIFICATES_FOLDER'])

    if not resolved_file_path.startswith(allowed_dir):
        print(f"ERRO: Tentativa de acesso a caminho inválido (fora da pasta permitida): {resolved_file_path}")
        return "Acesso negado: caminho inválido.", 403

    if os.path.exists(resolved_file_path) and os.path.isfile(resolved_file_path):
        try:
            return send_file(resolved_file_path, as_attachment=True)
        except Exception as e_send:
            print(f"ERRO CRÍTICO em send_file para '{resolved_file_path}': {e_send}")
            traceback.print_exc()
            return "Erro interno ao tentar servir o ficheiro.", 500
    else:
        print(f"ERRO: Ficheiro não encontrado para download: {resolved_file_path}")
        return "Arquivo não encontrado.", 404

if __name__ == '__main__':
    print("INFO: Iniciando servidor Flask em http://127.0.0.1:5000 (ou outro host/porta se configurado).")
    print("AVISO: Esta aplicação usa um sistema de autenticação simplificado (username como token e senhas hashed em JSON).")
    print("NÃO use para produção sem implementar autenticação e segurança robustas (ex: JWTs, OAuth2, gerenciamento de sessão adequado).")
    app.run(debug=True, host='0.0.0.0', port=5000)

